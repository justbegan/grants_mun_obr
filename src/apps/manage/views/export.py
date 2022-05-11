from typing import List
from django.db.models.fields import DateTimeField
from django.http import HttpResponseNotFound, HttpResponse
from django.shortcuts import render, get_object_or_404, redirect
from apps.project.models.organization import Organization
from apps.project.models.project import Project
from django_tables2 import RequestConfig
from docx import Document # https://python-docx.readthedocs.io/en/latest/
from docx.shared import Inches, Pt
from apps.expert.models import Expert


def example1():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')
#end example

def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)

def fdate(date):
    # return date.strftime("%d.%m.%Y %H:%M")
    return date.strftime("%d.%m.%Y")

def export_form1(request, projects:List[Project]):

    template_path = 'templates/export/form1.docx'
    document = Document(template_path)
    table = document.tables[0]
    
    """
    0. полное наименование организации
    1. ОГРН и (или) ИНН
    2. название и (или) краткое описание проекта (программы)
    3. дата начала и окончания срока реализации проекта (программы)
    4. форма предоставляемой поддержки (субсидия, грант в форме субсидии, грант)
    5. размер поддержки
    6. запрошенная организацией сумма поддержки
    """

    for project in projects:
        row_cells = table.add_row().cells
        # project: Project = projects[0]

        organization: Organization = project.organization
        row_cells[0].text = organization.full_name
        row_cells[1].text = 'ОГРН:' + organization.ogrn + '\n' + 'ИНН:' + organization.inn
        row_cells[2].text = project.title
        row_cells[3].text = fdate(project.start_date) + ' - ' + fdate(project.finish_date)
        row_cells[4].text = 'Грант'
        row_cells[5].text = str(project.budget_request_sum())
        row_cells[6].text = str(project.budget_request_sum())


    row = table.rows[1]
    remove_row(table, row)


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=form1.docx'
    document.save(response)

    return response


def export_form2(request, projects:List[Project]):

    template_path = 'templates/export/form2.docx'
    document = Document(template_path)
    table = document.tables[0]
    
    """
    0. полное наименование организации
    1. ОГРН и (или) ИНН
    2. название и (или) краткое описание проекта (программы)
    3. дата принятия решения об оказании поддержки
    4. Форма предоставленной поддержки (субсидия, грант в форме субсидии, грант)
    5. размер поддержки, указанный в решении о ее оказании
    6. размер поддержки, указанный в договоре (соглашении)
    7. общая сумма поддержки, фактически предоставленную организации
    8. сумма поддержки, перечисленную организации за счет гранта Президента Российской Федерации
    9. дата начала и окончания срока реализации проекта (программы), на осуществление которого предоставляется поддержка (срока, предусмотренного для использования поддержки)

    """

    for project in projects:
        row_cells = table.add_row().cells
        # project: Project = projects[0]

        organization: Organization = project.organization
        row_cells[0].text = organization.full_name
        row_cells[1].text = 'ОГРН:' + organization.ogrn + '\n' + 'ИНН:' + organization.inn
        row_cells[2].text = project.title
        row_cells[3].text = ''
        row_cells[4].text = 'Грант'
        row_cells[5].text = str(project.budget_request_sum())
        row_cells[6].text = str(project.budget_request_sum())
        row_cells[7].text = str(project.budget_request_sum())

        row_cells[8].text = ''
        row_cells[9].text = fdate(project.start_date) + ' - ' + fdate(project.finish_date)

    row = table.rows[1]
    remove_row(table, row)


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=form2.docx'
    document.save(response)

    return response



def export_form3(request, projects:List[Project]):

    template_path = 'templates/export/form3.docx'
    document = Document(template_path)
    table = document.tables[0]
    
    """
    0. № п/п
    1. № заявки
    2. Наименование организации
    3. Название проекта
    4. Общая сумма расходов по направлению, в руб.
    5. Запрашиваемая сумма
    6. Оценка эксперта 1
    7. Оценка эксперта 2
    8. Средняя оценка двух экспертов
    9. Балл за ОИПУ
    10. Балл за рейтинг Минэкономразвития
    11. Итоговая оценка по проекту
    """

    index = 0
    for project in projects:
        row_cells = table.add_row().cells
        # project: Project = projects[0]
        index += 1

        organization: Organization = project.organization
        
        row_cells[0].text = str(index)
        row_cells[1].text = str(project.id)
        row_cells[2].text = organization.full_name
        row_cells[3].text = project.title
        row_cells[4].text = str(project.budget_sum()).replace('.', ',')
        row_cells[5].text = str(project.budget_request_sum()).replace('.', ',')

        # if len(project.experts)>0:
        #     expert1: Expert = project.experts[0]
        #     if expert1:
        #         row_cells[6].text = expert1.

        row_cells[6].text = ''
        row_cells[7].text = ''
        i = 6
        for sheet in project.scoresheet_set.all():
            row_cells[i].text = str(round(sheet.score_total_sum(), 2)).replace('.', ',')
            i = i + 1
            if i > 7:
                break

        row_cells[8].text = str(project.averange_score()).replace('.', ',')
        row_cells[9].text = ''
        row_cells[10].text = ''
        row_cells[11].text = str(project.averange_score()).replace('.', ',')


    row = table.rows[1]
    remove_row(table, row)


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=form3.docx'
    document.save(response)

    return response

class RowForm4:

    all_count: int = 0

    DRAFT:int = 0 # Черновик
    NEW:int = 0 # Подана
    ON_CHECK:int = 0 # На проверке
    FIX:int = 0 # На доработку
    ON_EXAM:int = 0 # На экспертизе
    WIN:int = 0 # Победитель конкурса
    NOT_WIN:int = 0 # Проект, не получивший поддержку
    REJECT:int = 0 # Отклонена

    def __init__(self) -> None:
        self.all_count = 0

        self.DRAFT = 0
        self.NEW = 0
        self.ON_CHECK = 0
        self.FIX = 0
        self.ON_EXAM = 0
        self.WIN = 0
        self.NOT_WIN = 0
        self.REJECT = 0
        

def export_form4(request, projects:List[Project]):
    template_path = 'templates/export/form4.docx'
    document = Document(template_path)
    table = document.tables[0]
    
    """
    0. №
    1. Наименования района / гордского округа
    2. Общее количество поданных заявок
    3. Из них со статусом Черновик
    4. На проверке
    5. На доработку
    6. На экспертизе
    7. Отклонен
    8. Проект не получивший поддержку
    9. Победитель конкурса

    """

    _dict = {}

    for project in projects:
        
        # row_cells = table.add_row().cells
        # project: Project = projects[0]

        organization: Organization = project.organization

        munic = organization.geography_str()

        if not munic in _dict: 
            _dict[munic] = RowForm4()

        d: RowForm4 = _dict[munic]

        d.all_count += 1

        if project.status == project.DRAFT:
            d.DRAFT += 1
        elif project.status == project.NEW:
            d.NEW += 1
        elif project.status == project.ON_CHECK:
            d.ON_CHECK += 1
        elif project.status == project.FIX:
            d.FIX += 1
        elif project.status == project.ON_EXAM:
            d.ON_EXAM += 1
        elif project.status == project.WIN:
            d.WIN += 1
        elif project.status == project.NOT_WIN:
            d.NOT_WIN += 1
        elif project.status == project.REJECT:
            d.REJECT += 1


    index = 0
    for munic in _dict:
        index += 1
        d: RowForm4 = _dict[munic]
        row_cells = table.add_row().cells

        row_cells[0].text = str(index)
        row_cells[1].text = munic
        row_cells[2].text = str(d.all_count)

        row_cells[3].text = str(d.DRAFT)
        row_cells[4].text = str(d.ON_CHECK)
        row_cells[5].text = str(d.FIX)
        row_cells[6].text = str(d.ON_EXAM)
        row_cells[7].text = str(d.REJECT)
        row_cells[8].text = str(d.NOT_WIN)
        row_cells[9].text = str(d.WIN)


    row = table.rows[1]
    remove_row(table, row)

    # ИТОГО
    sum: RowForm4 = RowForm4()
    for munic in _dict:
        d: RowForm4 = _dict[munic]
        sum.all_count += d.all_count
        sum.DRAFT += d.DRAFT
        sum.NEW += d.NEW
        sum.ON_CHECK += d.ON_CHECK
        sum.FIX += d.FIX
        sum.ON_EXAM += d.ON_EXAM
        sum.WIN += d.WIN
        sum.NOT_WIN += d.NOT_WIN
        sum.REJECT += d.REJECT

    row_cells = table.add_row().cells
    row_cells[1].text = 'Итого'
    row_cells[2].text = str(sum.all_count)
    row_cells[3].text = str(sum.DRAFT)
    row_cells[4].text = str(sum.ON_CHECK)
    row_cells[5].text = str(sum.FIX)
    row_cells[6].text = str(sum.ON_EXAM)
    row_cells[7].text = str(sum.REJECT)
    row_cells[8].text = str(sum.NOT_WIN)
    row_cells[9].text = str(sum.WIN)
        



    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=form4.docx'
    document.save(response)

    return response


"""
-------Вопросы --------

Форма 1
? Откуда взять эти поля
4. форма предоставляемой поддержки (субсидия, грант в форме субсидии, грант)
5. размер поддержки


Форма 2
? Откуда взять эти поля
3. дата принятия решения об оказании поддержки
4. Форма предоставленной поддержки (субсидия, грант в форме субсидии, грант)
6. размер поддержки, указанный в договоре (соглашении)
7. общая сумма поддержки, фактически предоставленную организации
8. сумма поддержки, перечисленную организации за счет гранта Президента Российской Федерации

Форма 4
1. Наименования района / гордского округа - это собирается?

Тут такая логика экспертов есть?
6. Оценка эксперта 1
7. Оценка эксперта 2
8. Средняя оценка двух экспертов

9. Балл за ОИПУ
10. Балл за рейтинг Минэкономразвития
"""